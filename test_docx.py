from docx import Document
import io

try:
    doc = Document()
    doc.add_paragraph("Mission verification: AI can read Word documents.")
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    raw = buffer.getvalue()
    
    # Test reading
    test_doc = Document(io.BytesIO(raw))
    text = "\n".join([para.text for para in test_doc.paragraphs])
    print(f"Read result: {text}")
    print("STATUS: 100% OPERATIONAL")
except Exception as e:
    print(f"STATUS: FAILED - {e}")
