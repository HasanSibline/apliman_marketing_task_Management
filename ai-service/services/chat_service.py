import logging
from typing import Dict, List, Any, Optional
import google.generativeai as genai
from datetime import datetime
import json
import os
import aiohttp
import asyncio
from .context_learning import ContextLearningService

from config import get_config

logger = logging.getLogger(__name__)

class ChatService:
    """Service for handling conversational AI chat with context and memory"""

    def __init__(self, api_keys: List[str], provider: str = "gemini"):
        self.config = get_config()
        self.api_keys = api_keys if isinstance(api_keys, list) else [api_keys]
        self.current_key_index = 0
        self.api_key = self.api_keys[0] # For backwards compatibility with internal services
        self.provider = provider.lower()
        
        # Use the global config model if no environment override
        if self.provider == "groq":
            self.model_name = self.config.GROQ_MODEL
            self.base_url = "https://api.groq.com/openai/v1"
        else:
            self.model_name = self.config.GEMINI_MODEL
            self.base_url = "https://generativelanguage.googleapis.com/v1beta"
            
        self.learning_service = ContextLearningService(self.api_key, self.model_name)
        logger.info(f"✅ ChatService initialized with {self.provider} ({self.model_name}) and {len(self.api_keys)} API keys")

    def _rotate_api_key(self):
        """Rotate to the next API key"""
        if len(self.api_keys) > 1:
            old_index = self.current_key_index
            self.current_key_index = (self.current_key_index + 1) % len(self.api_keys)
            self.api_key = self.api_keys[self.current_key_index]
            # Update learning service with new key
            self.learning_service.api_key = self.api_key
            logger.warning(f"🔄 Rotating Chat API key from index {old_index} to {self.current_key_index}")
            return True
        return False

    async def process_chat_message(
        self,
        message: str,
        user_context: Dict[str, Any],
        user: Dict[str, Any],
        conversation_history: List[Dict[str, Any]],
        knowledge_sources: List[Dict[str, Any]],
        additional_context: Dict[str, Any] = None,
        is_deep_analysis: bool = False,
        company_name: str = None,
        files: Optional[List[Dict[str, Any]]] = None,
        user_token: Optional[str] = None
    ) -> Dict[str, Any]:
        """Process an incoming chat message with dynamic provider routing and multimodal support"""
        
        additional_context = additional_context or {}
        user_context = user_context or {}
        knowledge_sources = knowledge_sources or []
        files = files or []

        # Defensive flattening for double-nested lists (frequent frontend issue)
        if isinstance(user, list) and len(user) > 0:
            user = user[0]
        if isinstance(user_context, list) and len(user_context) > 0:
            user_context = user_context[0]
        if isinstance(additional_context, list) and len(additional_context) > 0:
            additional_context = additional_context[0]
        
        # Flatten knowledge_sources if it contains a list
        norm_ks = []
        for ks in knowledge_sources:
            if isinstance(ks, list):
                norm_ks.extend([x for x in ks if isinstance(x, dict)])
            elif isinstance(ks, dict):
                norm_ks.append(ks)
        knowledge_sources = norm_ks

        # Flatten files if double nested array received
        norm_files = []
        for f in files:
            if isinstance(f, list):
                norm_files.extend([x for x in f if isinstance(x, dict)])
            elif isinstance(f, dict):
                norm_files.append(f)
        files = norm_files

        if not isinstance(user, dict):
            user = {}

        # ── File-type flags (MUST be initialized before the try block regardless of whether files exist) ──
        has_media = False   # True if at least one image is attached
        has_docs = False    # True if at least one PDF/DOCX is attached
        document_text = ""  # Pre-extracted text from PDF/DOCX files

        logger.info(f"FILES RECEIVED: {len(files) if files else 'NONE'} files")
        if files:
            import json
            logger.info(f"FILES PAYLOAD (first item struct): {json.dumps(files[0], default=str)[:1000]}")
            
            # Determine if we have IMAGES or PDF/DOCX for multimodal/text extraction
            for file in files:
                mime = file.get("type", "")
                name = file.get("name", "")
                b64 = file.get("base64", "")
                
                is_image = any(mime.startswith(t) for t in ["image/"]) or name.lower().endswith((".jpg", ".jpeg", ".png", ".webp", ".gif"))
                is_pdf = name.lower().endswith(".pdf") or mime == "application/pdf"
                is_doc = name.lower().endswith((".docx", ".doc")) or "word" in mime.lower()

                if is_image:
                    has_media = True
                if is_pdf or is_doc:
                    has_docs = True

                # Extraction for all models (text context)
                if b64 and is_pdf:
                    try:
                        import base64 as b64_lib
                        import PyPDF2
                        import io
                        decoded = b64_lib.b64decode(b64)
                        reader = PyPDF2.PdfReader(io.BytesIO(decoded))
                        for page in reader.pages:
                            txt = page.extract_text()
                            if txt: document_text += txt + "\n"
                        logger.info(f"📄 Pre-extracted {len(document_text)} chars from PDF {name}")
                    except Exception as e:
                        logger.error(f"❌ PDF extraction failed: {str(e)}")
                elif b64 and is_doc:
                    try:
                        import base64 as b64_lib
                        from docx import Document
                        import io
                        decoded = b64_lib.b64decode(b64)
                        doc = Document(io.BytesIO(decoded))
                        text = "\n".join([para.text for para in doc.paragraphs])
                        if text: document_text += text + "\n"
                        logger.info(f"📄 Pre-extracted text from Word doc {name}")
                    except Exception as e:
                        logger.error(f"❌ DOCX extraction failed: {str(e)}")

            # Append document text to the user's message natively so any text provider can read it
            if document_text:
                document_text_header = "\n\n=== ATTACHED DOCUMENT CONTENT ===\n"
                message = f"{message}{document_text_header}{document_text[:30000]}"

        try:
            logger.info(f"Processing chat message (Files: {len(files) if files else 0}, HasMedia: {has_media}, HasDocs: {has_docs})")
            
            # Construct dynamic history block
            history_text = "CHAT HISTORY:\n"
            for msg in reversed(conversation_history[-7:]): # Last 7 for better context
                role_label = "ApliChat" if msg.get("role") == "assistant" else "User"
                history_text += f"{role_label}: {msg.get('content', '')}\n"

            # Create highly dynamic system prompt
            system_prompt = self._build_system_prompt(
                user=user,
                user_context=user_context,
                knowledge_sources=knowledge_sources,
                additional_context=additional_context,
                is_deep_analysis=is_deep_analysis,
                company_name=company_name,
                has_files=(has_media or has_docs) # Flag active if any asset is attached
            )

            # Generate response via appropriate provider
            if self.provider == "groq" and not has_media:
                # For Groq, we still use a flattened prompt for now (it has the appended document text!)
                full_prompt = f"{system_prompt}\n\n{history_text}\n\nUser: {message}\nApliChat:"
                response_text = await self._generate_via_groq(full_prompt)
            else:
                # Fallback mechanism: Groq lacks native multimodal vision endpoints.
                if self.provider == "groq" and has_media:
                    logger.warning("Visual files attached! Groq lacks vision. Falling back to Gemini.")
                    from config import get_config
                    config_instance = get_config()
                    self.model_name = config_instance.GEMINI_MODEL
                    self.base_url = "https://generativelanguage.googleapis.com/v1beta"
                    
                    fallback_keys = config_instance.get_api_keys()
                    if fallback_keys:
                        self.api_key = fallback_keys[0]
                        self.api_keys = fallback_keys
                    else:
                        raise Exception("Company uses Groq but attached a file. System requires GOOGLE_API_KEY environment variable for Gemini Vision Fallback!")

                if self.api_key and self.api_key.startswith("gsk_") and "generativelanguage" in self.base_url:
                    raise Exception("A Groq API key is being inappropriately sent to Google's Gemini endpoint. Please check system fallback keys.")

                # Primary attempt via REST (Gemini)
                try:
                    response_text = await self._generate_via_rest(
                        message=message,
                        system_prompt=system_prompt,
                        history_text=history_text,
                        files=files, 
                        user_token=user_token
                    )
                except Exception as rest_e:
                    error_msg = str(rest_e)
                    # CROSS-PROVIDER HOT SWAP: If Gemini entirely ran out of quota (429) and NO literal images are attached (has_media=False)
                    if "429" in error_msg and not has_media:
                        logger.critical("⚠️ Google Gemini hit a Hard 429 Rate Limit! Attempting violent cross-provider failover to Groq (Llama-3)...")
                        
                        # Dynamically aggressively extract Groq keys to save the response
                        import os
                        groq_keys = os.getenv("GROQ_API_KEYS", "").split(",")
                        if not groq_keys or not groq_keys[0].strip():
                            groq_keys = [os.getenv("GROQ_API_KEY", "")]
                            
                        groq_key = groq_keys[0].strip() if groq_keys and groq_keys[0].strip() else None
                        
                        if groq_key:
                            self.api_key = groq_key
                            self.model_name = os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile")
                            self.base_url = "https://api.groq.com/openai/v1"
                            logger.info(f"🔄 Cross-Provider Failover Initialized using {self.model_name}")
                            
                            full_prompt = f"{system_prompt}\n\n{history_text}\n\nUser: {message}\nApliChat:"
                            response_text = await self._generate_via_groq(full_prompt)
                        else:
                            raise rest_e # Give up if no Groq keys loaded natively
                    else:
                        raise rest_e # Propagate if not a 429 or requires vision
                
            response_text = response_text.strip()

            # Use AI to intelligently extract and update context
            learned_context = None
            try:
                learned_context = await self.learning_service.extract_and_update_context(
                    message=message,
                    existing_context=user_context,
                    conversation_history=conversation_history,
                    user_info=user
                )
            except Exception as learn_err:
                logger.warning(f"⚠️ Context learning failed (likely rate limited): {learn_err}")
                learned_context = None

            logger.info(f"✅ Generated chat response using {self.provider}")
            if learned_context:
                logger.debug(f"✅ Learned new context")

            return {
                "message": response_text,
                "contextUsed": True,
                "learnedContext": learned_context
            }

        except Exception as e:
            error_msg = str(e)
            logger.error(f"Error processing chat message: {error_msg}")
            
            # Surface the actual error message to the user/frontend for debugging
            detailed_msg = f"I encountered an error: {error_msg}"
            if "429" in error_msg:
                detailed_msg = "It looks like my AI quota just ran out. Please wait a minute and try again!"
            elif "API key was reported as leaked" in error_msg:
                detailed_msg = "Your AI API key has been revoked by the provider. Please update your company settings with a new key."
            
            return {
                "message": detailed_msg,
                "contextUsed": False,
                "learnedContext": None
            }

    async def _generate_via_groq(self, prompt: str) -> str:
        """Make a request to Groq API (OpenAI compatible)"""
        url = f"{self.base_url}/chat/completions"
        
        payload = {
            "model": self.model_name,
            "messages": [
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.5,
            "max_tokens": 2048,
            "stream": False
        }
        
        headers = {
            'Content-Type': 'application/json',
            'Authorization': f'Bearer {self.api_key}'
        }
        
        try:
            async with aiohttp.ClientSession() as session:
                async with session.post(url, headers=headers, json=payload, timeout=aiohttp.ClientTimeout(total=30)) as response:
                    if response.status == 200:
                        data = await response.json()
                        return data['choices'][0]['message']['content']
                    else:
                        error_text = await response.text()
                        logger.error(f"❌ Groq Chat API error ({response.status}): {error_text}")
                        raise Exception(f"Groq API failure ({response.status}): {error_text}")
        except Exception as e:
            logger.error(f"❌ Groq Chat request failed: {str(e)}")
            raise e

    async def _generate_via_rest(
        self, 
        message: str, 
        system_prompt: str,
        history_text: str,
        files: Optional[List[Dict[str, Any]]] = None, 
        user_token: Optional[str] = None
    ) -> str:
        """Make a request to Gemini API via REST with multi-modal parts"""
        attempts = 0
        last_error = None

        # --- Multimodal Support: Fetch assets and binary content ---
        media_parts = []
        text_content_parts = []
        file_count = 0

        if files:
            # Prefer absolute URLs sent by backend; fallback to robust local/remote guessing
            backend_base = os.getenv("BACKEND_URL", "").rstrip('/')
            
            for f in files:
                url = f.get("url", "")
                name = f.get("name", "file")
                mime = f.get("type", "")

                # --- STEP 1: PREFER EMBEDDED BASE64 (Eliminates Fetch Failures) ---
                embedded_b64 = f.get("base64")
                if embedded_b64:
                    logger.info(f"🚀 MULTIMODAL: Processing embedded Base64 for {name} ({mime})")
                    file_count += 1
                    is_image = any(mime.startswith(t) for t in ["image/"]) or name.lower().endswith((".jpg", ".jpeg", ".png", ".webp", ".gif"))
                    is_pdf = name.lower().endswith(".pdf") or mime == "application/pdf"
                    
                    if is_image:
                        actual_mime = mime if mime else "image/jpeg"
                        media_parts.append({
                            "inlineData": {
                                "mimeType": actual_mime,
                                "data": embedded_b64
                            }
                        })
                        logger.info(f"🖼️ Attached visual part via Base64: {name}")
                    elif is_pdf:
                        # Gemini 1.5/2.0 natively supports PDF parts!
                        media_parts.append({
                            "inlineData": {
                                "mimeType": "application/pdf",
                                "data": embedded_b64
                            }
                        })
                        logger.info(f"📄 Attached binary PDF part via Base64: {name}")
                    else:
                        # Fallback for Word/Text docs: they are already in the text message, but we can add them here too
                        try:
                            import base64 as b64_lib
                            decoded = b64_lib.b64decode(embedded_b64)
                            if name.lower().endswith((".docx", ".doc")):
                                # Word already handled by early extraction, adding a placeholder
                                text_content_parts.append({"text": f"[Analyzing Document: {name}]"})
                            else:
                                text = decoded.decode("utf-8", errors="replace")
                                text_content_parts.append({"text": f"[Attached File '{name}']: {text[:15000]}"})
                                logger.info(f"📄 Attached textual doc part via Base64: {name}")
                        except Exception as e:
                            logger.error(f"❌ Base64 decode error for {name}: {str(e)}")
                    continue 

                # --- STEP 2: FALLBACK TO URL FETCHING ---
                # Normalize URL for fetching
                full_url = url
                if not (url.startswith("http://") or url.startswith("https://")):
                    if backend_base:
                        url_sep = "" if url.startswith("/") else "/"
                        full_url = f"{backend_base}{url_sep}{url}"
                    else:
                        full_url = f"http://localhost:3001/{url.lstrip('/')}"
                
                try:
                    logger.info(f"✨ MULTIMODAL FETCH: Trying {name} from {full_url}")
                    headers = {}
                    if user_token:
                        headers['Authorization'] = user_token if user_token.startswith('Bearer ') else f'Bearer {user_token}'

                    async with aiohttp.ClientSession() as session:
                        async with session.get(full_url, headers=headers, timeout=aiohttp.ClientTimeout(total=20)) as file_res:
                            if file_res.status == 200:
                                raw = await file_res.read()
                                logger.info(f"✅ Downloaded {name} ({len(raw)} bytes)")
                                file_count += 1
                                
                                # 1. Binary Parts (Gemini Inline Data)
                                is_image = any(mime.startswith(t) for t in ["image/"]) or name.lower().endswith((".jpg", ".jpeg", ".png", ".webp", ".gif"))
                                if is_image:
                                    import base64
                                    actual_mime = mime if mime else "image/jpeg"
                                    b64 = base64.b64encode(raw).decode("utf-8")
                                    media_parts.append({
                                        "inlineData": {
                                            "mimeType": actual_mime,
                                            "data": b64
                                        }
                                    })
                                    logger.info(f"🖼️ Attached visual part via URL: {name}")
                                
                                # 2. Text Parts (Extracted content)
                                else:
                                    content = None
                                    if name.lower().endswith(".pdf"):
                                        import PyPDF2
                                        import io
                                        try:
                                            reader = PyPDF2.PdfReader(io.BytesIO(raw))
                                            pdf_text = ""
                                            for page in reader.pages:
                                                text_extract = page.extract_text()
                                                if text_extract:
                                                    pdf_text += text_extract + "\n"
                                            content = f"[Attached PDF '{name}']:\n{pdf_text[:25000]}"
                                            logger.info(f"📄 Extracted {len(pdf_text)} chars from downloaded PDF: {name}")
                                        except Exception as e:
                                            content = f"[Error reading PDF '{name}': {str(e)}]"
                                            logger.error(f"❌ PyPDF2 error for URL fetch {name}: {str(e)}")
                                    elif name.lower().endswith((".docx", ".doc")):
                                        try:
                                            from docx import Document
                                            import io
                                            doc = Document(io.BytesIO(raw))
                                            text = "\n".join([para.text for para in doc.paragraphs])
                                            content = f"[Attached Word Document '{name}']:\n{text[:25000]}"
                                        except Exception as e:
                                            content = f"[Error reading Word Doc '{name}': {str(e)}]"
                                    else:
                                        # Plain text, CSV, JSON, MD, etc.
                                        try:
                                            text = raw.decode("utf-8", errors="replace")
                                            content = f"[Attached File '{name}']:\n{text[:15000]}"
                                        except:
                                            content = f"[Attached Binary File '{name}' - Length: {len(raw)} bytes]"
                                    
                                    if content:
                                        text_content_parts.append({"text": content})
                                        logger.info(f"📄 Attached textual part: {name}")
                            else:
                                logger.error(f"❌ FETCH FAILED (Status {file_res.status}) for {full_url}")
                                text_content_parts.append({"text": f"(System Error: Could not retrieve file '{name}' for analysis. Status {file_res.status})"})
                except Exception as e:
                    logger.error(f"❌ MULTIMODAL EXCEPTION ({name}): {str(e)}")
                    text_content_parts.append({"text": f"(System Error: Failed to fetch file '{name}')"})

        # Final Prompt Construction: ATTACHMENTS FIRST, THEN RECENT HISTORY, THEN MESSAGE
        # This reordering is proven more effective for Gemini 1.5 context prioritization
        user_message_text = message if (message and message.strip()) else "(Analyzing attached assets...)"
        
        # Build prioritized parts list
        parts = []
        
        if media_parts:
            parts.extend(media_parts)
            logger.info(f"⚡ Vision: Attached {len(media_parts)} binary/media parts to prompt.")
            
        if text_content_parts:
            parts.extend(text_content_parts)
            logger.info(f"⚡ Documents: Attached {len(text_content_parts)} text/doc parts to prompt.")
            
        parts.append({
            "text": f"IMPORTANT: ANALYZE THE ABOVE ASSETS FIRST.\n\nRecent History:\n{history_text}\n\nUser Message: {user_message_text}"
        })

        if file_count > 0:
            parts.append({"text": f"\n[SYSTEM NOTICE: Task-specific analysis mode is ACTIVE for the {file_count} file(s) above. If the user asks about these files, ignore generic company knowledge and focus on the file content.]"})

        max_attempts = max(len(self.api_keys), 3)

        while attempts < max_attempts:
            current_key = self.api_key
            url = f"{self.base_url}/models/{self.model_name}:generateContent"
            
            payload = {
                "systemInstruction": {
                    "parts": [{"text": system_prompt}]
                },
                "contents": [
                    {
                        "role": "user",
                        "parts": parts
                    }
                ],
                "generationConfig": {
                    "temperature": 0.2, # Slightly lower for more reliable file analysis
                    "maxOutputTokens": 4096,
                    "topP": 0.9,
                }
            }

            for auth_method in ["query", "header"]:
                auth_url = f"{url}?key={current_key}" if auth_method == "query" else url
                try:
                    async with aiohttp.ClientSession() as session:
                        headers = {'Content-Type': 'application/json'}
                        if auth_method == "header":
                            headers['X-goog-api-key'] = current_key

                        async with session.post(auth_url, headers=headers, json=payload, timeout=aiohttp.ClientTimeout(total=50)) as response:
                            if response.status == 200:
                                data = await response.json()
                                if data.get('candidates'):
                                    candidate = data['candidates'][0]
                                    if candidate.get('content'):
                                        return candidate['content']['parts'][0]['text']
                                    elif candidate.get('finishReason'):
                                        return f"⚠️ Google Gemini chose not to respond due to Safety/Policy settings (Finish Reason: {candidate['finishReason']})"
                            
                            # Handle errors
                            try:
                                error_text = await response.text()
                            except Exception:
                                error_text = "Unknown Error"
                                
                            if response.status == 429:
                                logger.warning(f"Rate limited (429). Attempt {attempts+1}. Rotating keys...")
                                last_error = "Google Gemini Rate Limit Exceeded (429). Please wait a minute and try again."
                                break
                                
                            logger.warning(f"API Attempt failed ({response.status}): {error_text[:250]}")
                            last_error = error_text
                except Exception as e:
                    last_error = str(e)
                    continue
            
            attempts += 1
            if not self._rotate_api_key() and attempts < max_attempts:
                await asyncio.sleep(2)  # Wait longer for rate limits
                continue
            elif attempts < max_attempts:
                continue

        raise Exception(f"AI Multimodal analysis failed: {last_error}")


    def _build_system_prompt(
        self,
        user: Dict[str, Any],
        user_context: Dict[str, Any],
        knowledge_sources: List[Dict[str, Any]],
        additional_context: Dict[str, Any],
        is_deep_analysis: bool,
        company_name: str = None,
        has_files: bool = False
    ) -> str:
        """Build the system prompt with vision identity preservation"""
        
        response_style = "comprehensive" if is_deep_analysis else "conversational"
        company_name = company_name or "the company"
        
        prompt = f"""You are ApliChat, a state-of-the-art MULTIMODAL AI assistant for {company_name}. 

IMPORTANT IDENTITY GUIDELINES:
- You are a VISION-CAPABLE model (Gemini 2.0 Flash). 
- If files or images are attached, you CAN see and analyze them perfectly using your multimodal capabilities.
- NEVER say "I am a text-based model" or "I cannot see images/files". If a file is attached, analyze it deeply!
- If a file is attached, it takes ABSOLUTE PRIORITY. Provide a helpful summary or description immediately.
- You can read PDFs, images, and Word documents.

Your capabilities:
- Multimodal analysis of images, PDFs, and documents
- Answering questions about {company_name}'s business and operations
- Task and lifecycle management assistance
- Deep analytical reasoning

Current User Context:
- Name: {user.get('name', 'Analyst')}
- Role: {user.get('role', 'User')}
- Department: {(user.get('department') or {}).get('name', 'N/A')}

=== ANALYST MODE ===
{"[MODE: ACTIVE] Multiple files/images have been attached to this message. Review them immediately." if has_files else "[MODE: STANDARD] No new files attached."}

"""


        # Add user context (memory from past conversations)
        if user_context and len(user_context) > 0:
            prompt += "\nWhat I remember about you:\n"
            for key, value in user_context.items():
                if key != 'lastUpdated' and value:
                    prompt += f"- {key}: {value}\n"

        # Add company knowledge (COMPANY or OWN_COMPANY type)
        company_sources = [s for s in knowledge_sources if s.get('type') in ['COMPANY', 'OWN_COMPANY']]
        has_company_content = False
        
        if company_sources:
            prompt += f"\n=== About {company_name} ===\n"
            for source in company_sources[:3]:  # Limit to top 3
                if source.get('content'):
                    content = source['content'][:2000]
                    prompt += f"\n{source.get('name', 'Source')}:\n{content}\n"
                    has_company_content = True
                elif source.get('description'):
                    prompt += f"\n{source.get('name', 'Source')}: {source.get('description')}\n"
                    has_company_content = True
        
        # Warn if no company knowledge available
        if not has_company_content:
            prompt += f"\n⚠️ WARNING: No knowledge sources with content available for {company_name}. If asked about {company_name}, you MUST say you don't have information yet.\n"

        # Add competitor knowledge with competitive analysis instructions
        competitor_sources = [s for s in knowledge_sources if s.get('type') == 'COMPETITOR']
        has_competitor_content = False
        
        if competitor_sources:
            prompt += "\n=== COMPETITIVE INTELLIGENCE ===\n"
            prompt += f"Use this information to help {company_name} compete effectively:\n\n"
            
            for source in competitor_sources[:2]:  # Limit to top 2
                if source.get('content'):
                    content = source['content'][:1500]
                    prompt += f"\n{source.get('name', 'Competitor')}:\n{content}\n"
                    has_competitor_content = True
                elif source.get('description'):
                    prompt += f"\n{source.get('name', 'Competitor')}: {source.get('description')}\n"
                    has_competitor_content = True
            
            if has_competitor_content:
                prompt += f"""
COMPETITIVE STRATEGY GUIDELINES:
- When discussing competitors, identify {company_name}'s unique advantages and differentiators
- Suggest ways {company_name} can improve based on competitor strengths
- Highlight gaps in competitor offerings that {company_name} can exploit
- Recommend strategies to position {company_name} ahead of competitors
- Focus on value propositions that set {company_name} apart
- Never directly attack or disparage competitors - focus on {company_name}'s strengths
"""

        # Add task and ticket references
        if additional_context.get('referencedTasks'):
            prompt += "\n=== Specifically Referenced Tasks ===\n"
            for task in additional_context['referencedTasks']:
                prompt += f"""
Task: {task['title']} (TSK-{task.get('taskNumber', 'N/A')})
Status: {task.get('currentPhase', {}).get('name', 'Unknown')}
Priority: {task.get('priority', 'N/A')}
Due: {task.get('dueDate', 'No due date')}
"""

        if additional_context.get('referencedTickets'):
            prompt += "\n=== Specifically Referenced Tickets ===\n"
            for ticket in additional_context['referencedTickets']:
                requester_name = (ticket.get('requester') or {}).get('name', 'Unknown')
                target_dept = (ticket.get('receiverDept') or {}).get('name', 'Unknown')
                assignee_name = (ticket.get('assignee') or {}).get('name', 'Unassigned')
                
                comments_text = ""
                if ticket.get('comments'):
                    comments_text = "\nComments:\n"
                    for comment in ticket['comments']:
                        user_name = (comment.get('user') or {}).get('name', 'User')
                        comments_text += f"- {user_name}: {comment.get('comment', '')}\n"

                prompt += f"""
Ticket: {ticket.get('title', 'Untitled')} ({ticket.get('ticketNumber', 'N/A')})
Status: {ticket.get('status', 'Unknown')}
Priority: {ticket.get('priority', 'N/A')}
Requester: {requester_name}
Logistical Target: {target_dept}
Assignee: {assignee_name}
Description: {ticket.get('description', 'No description')}
{comments_text}
"""

        # User's Active Tickets Context
        if additional_context.get('userActiveTickets'):
            prompt += "\n=== User's Active Tickets ===\n"
            for t in additional_context['userActiveTickets']:
                prompt += f"- Ticket: {t.get('title')} ({t.get('ticketNumber')}) | Status: {t.get('status')} | Priority: {t.get('priority', 'N/A')}\n"
        
        # Give AI the user's workload context natively
        if additional_context.get('userAnalytics'):
            analytics = additional_context['userAnalytics']
            prompt += f"\n=== Current User Statistics ===\n"
            prompt += f"- Total Active Tasks: {analytics.get('activeTaskCount', 0)}\n"
            prompt += f"- Total Completed Tasks: {analytics.get('completedTaskCount', 0)}\n"

        if additional_context.get('userActiveTasks'):
            prompt += "\n=== User's Active Tasks ===\n"
            for task in additional_context['userActiveTasks']:
                phase_name = (task.get('currentPhase') or {}).get('name', 'Unknown')
                prompt += f"- Task: {task.get('title', 'Untitled')} | Phase: {phase_name} | Priority: {task.get('priority', 'N/A')} | Due: {task.get('dueDate', 'No due date')}\n"

        if additional_context.get('companyObjectives'):
            prompt += "\n=== Active Company Objectives ===\n"
            for obj in additional_context['companyObjectives']:
                prompt += f"- Objective: {obj.get('title')} | Status: {obj.get('status')}\n"

        if additional_context.get('companyQuarters'):
            prompt += "\n=== Active/Upcoming Quarters ===\n"
            for q in additional_context['companyQuarters']:
                prompt += f"- Quarter: {q.get('name')} {q.get('year')} | Status: {q.get('status')} | Period: {q.get('startDate')} to {q.get('endDate')}\n"

        # Add recent Microsoft meeting transcripts
        if additional_context.get('recentMeetings'):
            prompt += "\n=== Recent Microsoft Teams Meetings ===\n"
            prompt += "Use these transcripts to answer questions about recent discussions or decisions.\n"
            for meeting in additional_context['recentMeetings']:
                prompt += f"Meeting: {meeting.get('title')} ({meeting.get('date')})\n"
                prompt += f"Transcript (up to 4000 chars):\n{meeting.get('transcript')[:4000]}\n---\n"

        # Add mentioned users
        if additional_context.get('mentionedUsers'):
            prompt += "\n=== Mentioned Users ===\n"
            for u in additional_context['mentionedUsers']:
                task_count = len(u.get('assignedTasks', []))
                prompt += f"""
User: {u['name']} ({u.get('role', 'Unknown')})
Position: {u.get('position', 'Not specified')}
Active tasks: {task_count}
"""

        prompt += f"""

CRITICAL INSTRUCTIONS - FOLLOW STRICTLY:

1. KNOWLEDGE SOURCE PRIORITY:
   - When asked about "{company_name}", ONLY use information from the knowledge sources above
   - If knowledge sources have NO content about {company_name}, say: "I don't have detailed information about {company_name} yet. Please ask your administrator to add knowledge sources."
   - NEVER make up or invent information about {company_name}
   - NEVER guess what {company_name} does or what services they offer

2. COMPANY VS PLATFORM:
   - {company_name} is a BUSINESS/ORGANIZATION (use knowledge sources)
   - The task management platform is just a TOOL they're using
   - Don't confuse the two

3. GENERAL KNOWLEDGE:
   - For general questions (weather, facts, etc.), use your knowledge but be honest about limitations
   - If you cannot access real-time information, say so clearly
   - Don't pretend to search online if you can't

4. RESPONSE RULES:
   - Keep responses short and friendly unless asked for details
   - If you don't know something, say "I don't know" or "I don't have that information"
   - Never hallucinate or make up plausible-sounding but false information
   - When discussing competitors, only use information from competitor knowledge sources

5. TASK MANAGEMENT HELP:
   - For platform questions (e.g., "how do I create a task"), provide helpful guidance
   - Use task references and user mentions when available

"""

        return prompt

    def _build_conversation_history(self, history: List[Dict[str, Any]]) -> str:
        """Build conversation history text"""
        if not history or len(history) <= 1:
            return ""

        history_text = "\n=== Recent Conversation ===\n"
        for msg in history[-10:]:  # Last 10 messages
            role = "You" if msg['role'] == 'user' else "ApliChat"
            content = msg['content'][:200]  # Truncate long messages
            history_text += f"{role}: {content}\n"

        return history_text

    async def learn_from_task_history(
        self,
        user_context: Dict[str, Any],
        completed_tasks: List[Dict[str, Any]],
        active_tasks: List[Dict[str, Any]]
    ) -> Optional[Dict[str, Any]]:
        """
        Learn from user's task history to understand their work patterns
        
        Args:
            user_context: Existing user context
            completed_tasks: Recently completed tasks
            active_tasks: Currently active tasks
            
        Returns:
            Learned context from tasks
        """
        try:
            return await self.learning_service.learn_from_tasks(
                completed_tasks=completed_tasks,
                active_tasks=active_tasks,
                existing_context=user_context
            )
        except Exception as e:
            logger.error(f"Error learning from task history: {e}")
            return None
    
    async def learn_about_domain_interests(
        self,
        domain_topic: str,
        user_questions: List[str],
        existing_knowledge: Dict[str, Any]
    ) -> Optional[Dict[str, Any]]:
        """
        Learn what user wants to know about specific domains
        
        Args:
            domain_topic: Domain name (e.g., "company", "competitors")
            user_questions: Questions user asked about this domain
            existing_knowledge: Current knowledge about user's interests
            
        Returns:
            Updated knowledge about user's domain interests
        """
        try:
            return await self.learning_service.learn_about_domain(
                domain_topic=domain_topic,
                user_questions=user_questions,
                existing_knowledge=existing_knowledge
            )
        except Exception as e:
            logger.error(f"Error learning domain interests: {e}")
            return None

